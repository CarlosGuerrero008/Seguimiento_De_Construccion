import docx
import sys

doc_path = r'C:\Users\USUARIO\Seguimiento_De_Construccion\INFORME MENSUAL DE OBRA 9 OK.docx'
doc = docx.Document(doc_path)

# Extraer todo el texto
for paragraph in doc.paragraphs:
    print(paragraph.text)

# También extraer texto de tablas si las hay
print("\n\n=== TABLAS (si existen) ===\n")
for table in doc.tables:
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            row_text.append(cell.text)
        print(" | ".join(row_text))
    print("---")
